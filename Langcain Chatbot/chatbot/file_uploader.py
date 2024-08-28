from langchain.schema import Document
from typing import List
import PyPDF2
import docx

def load_files(file_paths: List[str]) -> List[Document]:
    documents = []
    for path in file_paths:
        if path.endswith('.pdf'):
            pdf_file_obj = open(path, 'rb')
            pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
            num_pages = pdf_reader.numPages
            text = ''
            for page in range(num_pages):
                page_obj = pdf_reader.getPage(page)
                text += page_obj.extractText()
            pdf_file_obj.close()
        elif path.endswith('.docx'):
            doc = docx.Document(path)
            text = ''
            for para in doc.paragraphs:
                text += para.text
        else:
            raise ValueError(f"Unsupported file type: {path}")
        
        documents.append(Document(page_content=text, metadata={"source": path}))
    return documents