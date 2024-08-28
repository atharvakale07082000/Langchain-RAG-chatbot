from pymongo import MongoClient
from datetime import datetime
from langchain_mongodb import MongoDBChatMessageHistory
from chatbot.vector_chroma import get_vector_store
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.schema import Document
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_huggingface import HuggingFaceEndpoint
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
import io
import PyPDF2
import docx

class Chatbot:
    def __init__(self, session_id: str):
        self.session_id = session_id
        self.memory = self.get_chat_memory(session_id)
        self.vector_store = get_vector_store()
        self.api_token = "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"

        # MongoDB setup
        self.mongo_client = MongoClient("mongodb://localhost:27017")  # Replace with your MongoDB connection string
        self.db = self.mongo_client["chatbot_db"]
        self.chat_collection = self.db["chat_history"]

        # Define a prompt template
        self.prompt_template = PromptTemplate(
            input_variables=["input", "context"],
            template="Given the context: {context}, answer the question: {input}"
        )

        # Initialize the HuggingFace Endpoint
        self.llm = HuggingFaceEndpoint(
            repo_id="OpenAssistant/oasst-sft-4-pythia-12b-epoch-3.5",  # Replace with your preferred model
            huggingfacehub_api_token=self.api_token,
            max_new_tokens=128  # Adjust as needed
        )

        # Initialize the history-aware retriever
        self.retriever = create_history_aware_retriever(
            retriever=self.vector_store.as_retriever(),
            prompt=self.prompt_template,
            llm=self.llm
        )

    def get_chat_memory(self, session_id: str):
        connection_string = "mongodb://localhost:27017"  # Replace with your MongoDB connection string
        return MongoDBChatMessageHistory(
            connection_string=connection_string,
            database_name="chatbot_db",
            collection_name="chat_memory",
            session_id=session_id
        )

    def load_documents(self, filename: str, content: bytes):
        file_like_object = io.BytesIO(content)

        if filename.endswith('.pdf'):
            text = self.extract_pdf_text(file_like_object)
        elif filename.endswith('.docx'):
            text = self.extract_docx_text(file_like_object)
        else:
            raise ValueError(f"Unsupported file type: {filename}")

        # Use recursive text splitter to split text into chunks
        chunks = self.split_text_into_chunks(text)

        # Add chunks to vector store
        for i, chunk in enumerate(chunks):
            document = Document(page_content=chunk, metadata={"source": f"{filename}_chunk_{i}"})
            print(f"Adding document chunk to vector store: {filename}_chunk_{i}")
            self.vector_store.add_documents([document])

    def extract_pdf_text(self, file_like_object: io.BytesIO) -> str:
        pdf_reader = PyPDF2.PdfReader(file_like_object)
        text = ''.join(page.extract_text() for page in pdf_reader.pages)
        return text

    def extract_docx_text(self, file_like_object: io.BytesIO) -> str:
        doc = docx.Document(file_like_object)
        text = ''.join(para.text for para in doc.paragraphs)
        return text

    def split_text_into_chunks(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> list:
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,      # Adjust the chunk size as needed
            chunk_overlap=150  # Adjust the chunk overlap as needed
        )
        return text_splitter.split_text(text)

    def save_to_mongo(self, session_id: str, query: str, response: str):
        record = {
            "session_id": session_id,
            "query": query,
            "response": response,
            "timestamp": datetime.utcnow()
        }
        self.chat_collection.insert_one(record)

    def answer_query(self, query: str):
        # Load context messages if they are part of the memory
        try:
            context_messages = self.memory.messages  # Use the messages property to get all messages
            context = " ".join(message.content for message in context_messages)
        except Exception as e:
            print(f"Error loading chat memory: {e}")
            context = ""

        # Use the retriever to get relevant documents based on the query and context
        try:
            # Create a custom chain that combines the retriever and llm
            combine_docs_chain = create_stuff_documents_chain(self.llm, self.prompt_template)

            # Create a retrieval chain with the combined document chain
            chain = create_retrieval_chain(
                retriever=self.retriever,
                combine_docs_chain=combine_docs_chain
            )

            # Run the chain to get the answer
            docs = self.retriever.invoke({"input": query, "context": context})
            doc_contents = [doc.page_content for doc in docs]
            answer = combine_docs_chain.invoke({"input": query, "context": context, "documents": doc_contents})

            # Save the user query and AI response to the chat memory
            self.save_to_mongo(self.session_id, query, answer)

            return answer
        except Exception as e:
            print(f"Error processing query: {e}")
            return "An error occurred while processing your query."
